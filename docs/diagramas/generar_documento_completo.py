#!/usr/bin/env python3
"""
generar_documento_completo.py
Genera documentacion_completa.docx combinando:
  - Especificacion de Casos de Uso (actores + CU por modulo)
  - Diagramas de Casos de Uso (imagenes PNG)
  - Diagramas Entidad-Relacion por modulo (imagenes PNG)
  - Diagramas de Secuencia (imagenes PNG)
  - Diagrama de Arquitectura y Requerimientos
Uso: python generar_documento_completo.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
GREEN  = RGBColor(0x27, 0xAE, 0x60)
DARK   = RGBColor(0x2C, 0x3E, 0x50)
GRAY   = RGBColor(0x66, 0x66, 0x66)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

DOCS = Path(r"d:\tita2026\docs\diagramas")
IMGS = DOCS / "diagrams_png"
OUT  = DOCS / "documentacion_completa.docx"


# ── OOXML helpers ─────────────────────────────────────────────────────────────
def cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def border_bottom(para, color="27AE60", sz=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), color)
    pBdr.append(bot); pPr.append(pBdr)

def border_left(para, color="27AE60", sz=18):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), "4"); left.set(qn("w:color"), color)
    pBdr.append(left); pPr.append(pBdr)

def set_repeat_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


# ── Doc setup ─────────────────────────────────────────────────────────────────
def make_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)
    sec.top_margin = Inches(0.9); sec.bottom_margin = Inches(0.9)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    return doc


# ── Texto helpers ─────────────────────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_heading(text, 1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.color.rgb = NAVY; r.font.size = Pt(16)
    border_bottom(p, "1E3A5F", 18)
    return p

def h2(doc, text):
    p = doc.add_heading(text, 2)
    for r in p.runs:
        r.font.color.rgb = NAVY; r.font.size = Pt(13)
    border_bottom(p, "27AE60", 10)
    return p

def h3(doc, text):
    p = doc.add_heading(text, 3)
    for r in p.runs:
        r.font.color.rgb = DARK; r.font.size = Pt(11)
    return p

def body(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text); r.font.size = Pt(10.5)
    return p

def quote(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.3); p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    for r in p.runs:
        r.font.size = Pt(10.5); r.font.italic = True; r.font.color.rgb = DARK
    border_left(p)
    return p

def label_val(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{label}: ")
    r1.font.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
    r2 = p.add_run(value); r2.font.size = Pt(10.5)

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(10)
    return p

def diagram_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(8.5); r.font.italic = True; r.font.color.rgb = GRAY

def section_divider(doc, title, subtitle=""):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = NAVY
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(11); r2.font.color.rgb = GREEN; r2.font.bold = True
    doc.add_paragraph()


# ── Tablas ────────────────────────────────────────────────────────────────────
def table_actor(doc, actors):
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    set_repeat_header(hdr)
    for i, txt in enumerate(["ID", "Actor", "Descripcion", "Nivel de acceso"]):
        c = hdr.cells[i]; c.text = txt
        cell_bg(c, "1E3A5F")
        for r in c.paragraphs[0].runs:
            r.font.color.rgb = WHITE; r.font.bold = True; r.font.size = Pt(9)
    for ri, (aid, nombre, desc, acceso) in enumerate(actors):
        row = tbl.add_row().cells
        row[0].text = aid; row[1].text = nombre
        row[2].text = desc; row[3].text = acceso
        bg = "E8F5E9" if ri % 2 == 0 else "FFFFFF"
        for ci in range(4):
            cell_bg(row[ci], bg)
            for r in row[ci].paragraphs[0].runs: r.font.size = Pt(9)
        for r in row[0].paragraphs[0].runs: r.font.bold = True; r.font.color.rgb = NAVY
        for r in row[1].paragraphs[0].runs: r.font.bold = True
    doc.add_paragraph()

def table_cu(doc, casos):
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    set_repeat_header(hdr)
    for i, txt in enumerate(["ID", "Caso de Uso", "Actor(es)", "Descripcion"]):
        c = hdr.cells[i]; c.text = txt
        cell_bg(c, "1E3A5F")
        for r in c.paragraphs[0].runs:
            r.font.color.rgb = WHITE; r.font.bold = True; r.font.size = Pt(9)
    for ri, (cid, nombre, actores, desc) in enumerate(casos):
        row = tbl.add_row().cells
        row[0].text = cid; row[1].text = nombre
        row[2].text = actores; row[3].text = desc
        bg = "F0F7F4" if ri % 2 == 0 else "FFFFFF"
        for ci in range(4):
            cell_bg(row[ci], bg)
            for r in row[ci].paragraphs[0].runs: r.font.size = Pt(9)
        for r in row[0].paragraphs[0].runs:
            r.font.bold = True; r.font.color.rgb = GREEN
        for r in row[1].paragraphs[0].runs: r.font.bold = True
    doc.add_paragraph()


def img(doc, png_name, width=6.3, caption=""):
    path = IMGS / png_name
    if not path.exists():
        body(doc, f"[Imagen no disponible: {png_name}]")
        return
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        if caption:
            diagram_caption(doc, caption)
    except Exception as e:
        body(doc, f"[Error al insertar imagen {png_name}: {e}]")


# ════════════════════════════════════════════════════════════════════════════
#   DATOS DEL DOCUMENTO
# ════════════════════════════════════════════════════════════════════════════

ACTORES = [
    ("A-01", "Administrador",
     "Usuario con acceso total al sistema. Gestiona configuracion general, usuarios, "
     "reportes avanzados, facturacion y parametros del negocio.",
     "Acceso completo (CRUD en todos los modulos)"),
    ("A-02", "Cajero",
     "Operador del punto de venta. Registra ventas de cantina y almuerzo, procesa "
     "pagos, abre y cierra caja y gestiona tarjetas prepagas en el POS.",
     "Acceso operativo: POS, caja, tarjetas, comprobantes"),
    ("A-03", "Padre / Responsable",
     "Cliente externo que accede al portal para consultar el saldo prepago de su hijo, "
     "configurar topes de gasto diarios y recibir notificaciones automaticas.",
     "Acceso restringido: portal web, consulta y configuracion propia"),
    ("A-04", "Sistema",
     "Actor secundario automatizado. Ejecuta triggers de base de datos, actualiza stock, "
     "genera alertas de saldo bajo, registra bitacoras de auditoria y calcula IVA.",
     "Acciones internas automatizadas, sin intervencion humana"),
]

MODULOS = [
    {
        "num": "1", "titulo": "Gestion de Clientes",
        "objetivo": (
            "Disenar e implementar un modulo de gestion de clientes que incorpore la "
            "vinculacion directa cliente-hijo y el control de cuentas corrientes prepagas, "
            "permitiendo a los padres configurar topes de gasto diarios y recibir alertas "
            "automaticas de saldo bajo."
        ),
        "actores": "Administrador, Padre / Responsable, Sistema",
        "cu_diag":  "diag_06.png",  # CU-02 Clientes
        "er_diag":  "diag_18.png",  # ER-02 Clientes
        "er_label": "ER-02 — Modulo Clientes e Hijos",
        "er_desc": (
            "Clientes (padres/responsables) con sus hijos/alumnos, cuenta corriente prepaga, "
            "historial de grados escolares y restricciones alimenticias por hijo."
        ),
        "casos": [
            ("CU-01", "Registrar cliente", "Administrador",
             "El administrador ingresa los datos del responsable (nombre, RUC/CI, contacto) y crea el perfil."),
            ("CU-02", "Vincular hijo a cliente", "Administrador",
             "Se asocia uno o mas hijos al cliente responsable estableciendo la relacion de dependencia."),
            ("CU-03", "Configurar tope de gasto diario", "Padre / Responsable",
             "El padre define el limite maximo de gasto diario que puede efectuar su hijo en el POS."),
            ("CU-04", "Consultar cuenta corriente", "Padre / Responsable, Administrador",
             "Se visualiza el historial de recargas, consumos y saldo disponible de la cuenta prepaga del hijo."),
            ("CU-05", "Cargar saldo prepago", "Administrador, Cajero",
             "El operador acredita un monto en la cuenta corriente del hijo mediante efectivo u otro medio de pago."),
            ("CU-06", "Recibir alerta de saldo bajo", "Sistema, Padre / Responsable",
             "El sistema detecta que el saldo disponible cae por debajo del umbral configurado y notifica al padre."),
            ("CU-07", "Consultar historial de consumos", "Padre / Responsable, Administrador",
             "Lista cronologica de todas las transacciones realizadas por el hijo en el punto de venta."),
        ],
    },
    {
        "num": "2", "titulo": "Administracion de Proveedores",
        "objetivo": (
            "Desarrollar un modulo de administracion de proveedores que gestione de manera "
            "totalmente aislada e independiente sus respectivas cuentas corrientes (cuentas "
            "por pagar), garantizando la separacion estricta con respecto a los saldos de los clientes."
        ),
        "actores": "Administrador",
        "cu_diag":  "diag_12.png",  # CU-08 Compras (contiene proveedores)
        "er_diag":  "diag_24.png",  # ER-08 Compras/Proveedores
        "er_label": "ER-08 — Modulo Compras y Proveedores",
        "er_desc": (
            "Compras a proveedores con detalle por producto, pagos con aplicacion a multiples "
            "compras, notas de credito y cuenta corriente de proveedor completamente aislada."
        ),
        "casos": [
            ("CU-08", "Registrar proveedor", "Administrador",
             "El administrador crea el perfil del proveedor con razon social, RUC, contacto y condiciones de pago."),
            ("CU-09", "Consultar cuenta corriente de proveedor", "Administrador",
             "Se visualiza el saldo de la deuda pendiente con el proveedor, separado de las cuentas de clientes."),
            ("CU-10", "Registrar deuda por compra", "Administrador, Sistema",
             "Al procesar una compra, el sistema acredita automaticamente la deuda en la cuenta del proveedor."),
            ("CU-11", "Registrar pago a proveedor", "Administrador",
             "El administrador registra el pago total o parcial de una deuda, actualizando el saldo por pagar."),
            ("CU-12", "Consultar historial de transacciones", "Administrador",
             "Listado de todas las compras y pagos asociados a un proveedor especifico con fechas e importes."),
        ],
    },
    {
        "num": "3", "titulo": "Arqueo y Cierre de Caja Diario",
        "objetivo": (
            "Disenar e implementar un modulo de arqueo y cierre de caja diario que registre "
            "de forma automatizada los ingresos por multiples medios de pago, permitiendo "
            "conciliar el efectivo fisico con las transacciones del sistema y documentar "
            "cualquier desvio financiero para su posterior auditoria."
        ),
        "actores": "Cajero, Administrador, Sistema",
        "cu_diag":  "diag_13.png",  # CU-09 Contabilidad y Caja
        "er_diag":  "diag_25.png",  # ER-09 Contabilidad
        "er_label": "ER-09 — Modulo Contabilidad / Caja",
        "er_desc": (
            "Gestion de cajas fisicas, cierres de caja diarios con movimientos clasificados "
            "por medio de pago, facturas preimpresas legales y datos de la empresa emisora."
        ),
        "casos": [
            ("CU-13", "Abrir caja del dia", "Cajero, Administrador",
             "El operador registra el monto inicial en efectivo y habilita la caja para recibir transacciones."),
            ("CU-14", "Registrar ingreso por medio de pago", "Sistema",
             "Cada venta acredita automaticamente el monto en el registro de caja segun el medio de pago utilizado."),
            ("CU-15", "Registrar egreso / retiro", "Cajero, Administrador",
             "Se documenta cualquier salida de dinero de la caja con su justificacion."),
            ("CU-16", "Realizar arqueo fisico", "Cajero",
             "El cajero ingresa el total contado fisicamente en caja al momento del cierre."),
            ("CU-17", "Cerrar caja y calcular diferencia", "Cajero, Sistema",
             "El sistema calcula la diferencia entre el monto esperado y el contado fisico, generando el reporte."),
            ("CU-18", "Generar reporte de cierre de caja", "Administrador, Sistema",
             "El sistema produce el comprobante del cierre con detalle por medio de pago e ingresos/egresos."),
        ],
    },
    {
        "num": "4", "titulo": "Punto de Venta (POS)",
        "objetivo": (
            "Disenar una interfaz de Punto de Venta (POS) optimizada para la alta concurrencia "
            "de los recreos, que diferencie de manera agil las operaciones de 'almuerzo' y "
            "'cantina/recreo', con el fin de minimizar el tiempo medio de atencion por estudiante."
        ),
        "actores": "Cajero, Sistema",
        "cu_diag":  "diag_09.png",  # CU-05 Ventas
        "er_diag":  "diag_21.png",  # ER-05 Ventas
        "er_label": "ER-05 — Modulo Ventas / POS",
        "er_desc": (
            "Ventas con desglose de IVA, pagos multiples por venta, notas de credito y "
            "aplicacion de pagos a multiples ventas. Vinculacion con medios de pago y tarjetas."
        ),
        "casos": [
            ("CU-19", "Registrar venta de cantina / recreo", "Cajero",
             "El cajero selecciona productos del catalogo y procesa la venta en modalidad cantina."),
            ("CU-20", "Registrar consumo de almuerzo", "Cajero",
             "El sistema registra el consumo del plan de almuerzo y debita el importe de la cuenta prepaga."),
            ("CU-21", "Procesar pago con tarjeta prepaga", "Cajero, Sistema",
             "El sistema debita el monto del saldo disponible en la tarjeta, verificando topes y disponibilidad."),
            ("CU-22", "Procesar pago en efectivo u otro medio", "Cajero",
             "El cajero registra el pago con efectivo, tarjeta de debito/credito u otro medio habilitado."),
            ("CU-23", "Anular venta", "Cajero, Administrador",
             "Se revierte una venta registrada por error, devolviendo el saldo y ajustando el stock."),
            ("CU-24", "Emitir comprobante de venta", "Sistema",
             "El sistema genera el ticket o comprobante de la transaccion para entrega al cliente."),
            ("CU-25", "Consultar historial de ventas del turno", "Cajero, Administrador",
             "Listado de todas las transacciones procesadas durante el turno activo de la caja."),
        ],
    },
    {
        "num": "5", "titulo": "Compras",
        "objetivo": (
            "Construir un modulo de compras articulado con el inventario centralizado y el "
            "registro de proveedores, que permita procesar comprobantes de recepcion de insumos "
            "y actualizar de forma automatica los costos operativos de los productos."
        ),
        "actores": "Administrador, Sistema",
        "cu_diag":  "diag_12.png",  # CU-08 Compras
        "er_diag":  "diag_24.png",  # ER-08 Compras
        "er_label": "ER-08 — Modulo Compras (detalle de inventario)",
        "er_desc": (
            "Compras a proveedores con detalle por producto, actualizacion automatica de stock "
            "e inventario, pagos parciales y notas de credito de proveedor."
        ),
        "casos": [
            ("CU-26", "Crear orden de compra", "Administrador",
             "El administrador genera una solicitud formal de compra de insumos a un proveedor."),
            ("CU-27", "Recepcionar insumos", "Administrador",
             "Se registra la recepcion fisica de los insumos segun el comprobante del proveedor."),
            ("CU-28", "Actualizar stock automaticamente", "Sistema",
             "Al confirmar la recepcion, el sistema incrementa las unidades disponibles en el inventario."),
            ("CU-29", "Actualizar costo operativo del producto", "Sistema",
             "El sistema recalcula el costo de los productos en funcion del precio de la compra recepcionada."),
            ("CU-30", "Vincular compra a cuenta de proveedor", "Sistema",
             "La compra queda asociada al proveedor, generando automaticamente la deuda en su cuenta corriente."),
            ("CU-31", "Consultar historial de compras", "Administrador",
             "Listado de todas las compras realizadas, filtrable por proveedor, fecha o producto."),
        ],
    },
    {
        "num": "6", "titulo": "Facturacion Legal",
        "objetivo": (
            "Incorporar un motor de facturacion legal que aplique las normativas fiscales "
            "vigentes, validando el campo estructurado RUC_CI, automatizando el calculo de "
            "las tasas del IVA y garantizando la correlacion numerica exacta de las facturas preimpresas."
        ),
        "actores": "Administrador, Cajero, Sistema",
        "cu_diag":  "diag_13.png",  # CU-09 Contabilidad (incluye facturacion)
        "er_diag":  "diag_25.png",  # ER-09 Contabilidad (incluye Factura)
        "er_label": "ER-09 — Modulo Contabilidad / Facturacion Legal",
        "er_desc": (
            "Facturas preimpresas legales vinculadas a ventas o cargas de saldo, con calculo "
            "de IVA (10%/5%/exenta), validacion de RUC/CI y control de correlacion numerica."
        ),
        "casos": [
            ("CU-32", "Emitir factura legal", "Administrador, Cajero",
             "Se registra una factura preimpresa vinculada a una venta o carga de saldo."),
            ("CU-33", "Validar RUC / CI del cliente", "Sistema",
             "El sistema verifica el formato del RUC o CI antes de emitir la factura."),
            ("CU-34", "Calcular IVA automaticamente", "Sistema",
             "El motor fiscal calcula IVA 10%, IVA 5% y monto exento segun la naturaleza del bien o servicio."),
            ("CU-35", "Verificar correlacion numerica", "Sistema",
             "El sistema garantiza que el numero de factura no este duplicado y respete la secuencia del talonario."),
            ("CU-36", "Anular factura", "Administrador",
             "El administrador anula una factura por error, registrando el motivo y bloqueando el numero."),
            ("CU-37", "Consultar facturas emitidas", "Administrador",
             "Listado de todas las facturas del periodo, filtrable por numero, cliente o fecha."),
        ],
    },
    {
        "num": "7", "titulo": "Tarjetas y Saldo Prepago",
        "objetivo": (
            "Desarrollar un mecanismo de control para tarjetas de uso exclusivo de la cantina, "
            "que habilite el procesamiento rapido de recargas de saldo y el debito automatizado "
            "de transacciones en el punto de venta."
        ),
        "actores": "Administrador, Cajero, Padre / Responsable, Sistema",
        "cu_diag":  "diag_08.png",  # CU-04 Core/Tarjetas
        "er_diag":  "diag_20.png",  # ER-04 Core
        "er_label": "ER-04 — Modulo Core / Tarjetas y Saldo",
        "er_desc": (
            "Tarjetas monedero vinculadas a hijos, cargas de saldo, consumos, medios de pago, "
            "limites de transaccion por rol y registro de alertas de saldo bajo."
        ),
        "casos": [
            ("CU-38", "Activar tarjeta prepaga", "Administrador",
             "El administrador registra una nueva tarjeta, la vincula al hijo y la deja habilitada."),
            ("CU-39", "Recargar saldo en tarjeta", "Cajero, Administrador",
             "El operador acredita un importe en la tarjeta del estudiante, registrando el medio de pago."),
            ("CU-40", "Debitar saldo en transaccion POS", "Sistema",
             "Al procesar una venta, el sistema descuenta automaticamente el importe verificando disponibilidad y topes."),
            ("CU-41", "Bloquear tarjeta", "Administrador, Padre / Responsable",
             "Se inhabilita temporalmente una tarjeta por extravio, robo o solicitud del responsable."),
            ("CU-42", "Consultar saldo y movimientos", "Padre / Responsable, Administrador",
             "Se visualiza el saldo actual y el historial de recargas y consumos de la tarjeta."),
            ("CU-43", "Generar alerta de saldo insuficiente", "Sistema",
             "El sistema detecta que el saldo cae bajo el umbral configurado y notifica al padre."),
        ],
    },
    {
        "num": "8", "titulo": "Reportes Analiticos",
        "objetivo": (
            "Implementar un panel de reportes analiticos que consolide datos historicos de ventas "
            "e inventarios (filtrados por fecha, cajero, rotacion y rentabilidad), proveyendo a la "
            "administracion herramientas visuales para la toma de decisiones."
        ),
        "actores": "Administrador",
        "cu_diag":  "diag_07.png",  # CU-03 Productos (base de precios/analisis)
        "er_diag":  "diag_19.png",  # ER-03 Productos (precios, historico, categorias)
        "er_label": "ER-03 — Modulo Productos y Precios (base de reportes)",
        "er_desc": (
            "Catalogo de productos con categorias jerarquicas, multiples listas de precios, "
            "asignacion de impuestos e historico de precios. Base de los reportes de rentabilidad y rotacion."
        ),
        "casos": [
            ("CU-44", "Generar reporte de ventas por periodo", "Administrador",
             "El sistema consolida todas las ventas del rango de fechas seleccionado."),
            ("CU-45", "Filtrar ventas por cajero", "Administrador",
             "Se desglosan las ventas segun el operador que las proceso."),
            ("CU-46", "Consultar rotacion de productos", "Administrador",
             "El reporte muestra los productos mas y menos vendidos en el periodo."),
            ("CU-47", "Analizar rentabilidad por producto", "Administrador",
             "El sistema calcula margen de ganancia unitario comparando precio de venta con costo de compra."),
            ("CU-48", "Visualizar dashboard de indicadores", "Administrador",
             "Panel con metricas clave: ventas del dia, saldo en caja, stock critico y alertas activas."),
            ("CU-49", "Exportar reporte historico", "Administrador",
             "El administrador descarga el reporte seleccionado en formato Excel o PDF."),
        ],
    },
    {
        "num": "9", "titulo": "Arquitectura de Base de Datos e Integridad",
        "objetivo": (
            "Disenar la arquitectura de la base de datos incorporando triggers y restricciones "
            "de integridad atomica para automatizar el control de stock en las ventas, asegurar "
            "la consistencia financiera en los cierres de caja y alimentar bitacoras de auditoria interna."
        ),
        "actores": "Sistema, Administrador",
        "cu_diag":  "diag_11.png",  # CU-07 Inventario (triggers de stock)
        "er_diag":  "diag_23.png",  # ER-07 Inventario (lotes, ajustes, control de stock)
        "er_label": "ER-07 — Modulo Inventario (control de stock y triggers)",
        "er_desc": (
            "Stock por producto, ajustes con aprobacion de doble nivel, lotes con fecha de "
            "vencimiento, costos historicos y alertas automaticas de stock critico via triggers."
        ),
        "casos": [
            ("CU-50", "Ejecutar trigger de control de stock en ventas", "Sistema",
             "Al registrar una venta, el trigger descuenta automaticamente el inventario y bloquea si stock es insuficiente."),
            ("CU-51", "Validar integridad atomica en transacciones", "Sistema",
             "Cada operacion es completamente exitosa o completamente revertida, sin estados intermedios."),
            ("CU-52", "Actualizar cuentas corrientes en cierre de caja", "Sistema",
             "Los triggers consolidan los movimientos del dia en cuentas corrientes al cerrar la caja."),
            ("CU-53", "Registrar evento en bitacora de auditoria", "Sistema",
             "Cada accion critica queda registrada con usuario, fecha y hora para auditoria interna."),
            ("CU-54", "Consultar bitacora de auditoria", "Administrador",
             "El administrador accede al historial de acciones criticas para resolver discrepancias."),
            ("CU-55", "Gestionar restricciones de integridad referencial", "Sistema",
             "El sistema aplica claves foraneas y restricciones de unicidad para prevenir inconsistencias."),
        ],
    },
    {
        "num": "10", "titulo": "Control de Accesos y Usuarios",
        "objetivo": (
            "Implementar un modelo de control de accesos basado en roles jerarquicos "
            "(administrador, cajero, padre) que restrinja las funciones del sistema segun "
            "el perfil del usuario y registre las acciones criticas para la resolucion de discrepancias."
        ),
        "actores": "Administrador, Cajero, Padre / Responsable, Sistema",
        "cu_diag":  "diag_05.png",  # CU-01 Usuarios
        "er_diag":  "diag_17.png",  # ER-01 Usuarios
        "er_label": "ER-01 — Modulo Usuarios y Autenticacion",
        "er_desc": (
            "Usuarios del sistema con autenticacion JWT, roles y permisos RBAC, sesiones activas, "
            "intentos de login, recuperacion de contrasena y soporte para autenticacion de dos factores (2FA)."
        ),
        "casos": [
            ("CU-56", "Registrar usuario en el sistema", "Administrador",
             "El administrador crea la cuenta de un nuevo usuario asignando credenciales, rol y permisos."),
            ("CU-57", "Autenticar usuario con JWT", "Sistema",
             "El sistema valida credenciales y emite un token JWT firmado que habilita el acceso."),
            ("CU-58", "Asignar rol jerarquico", "Administrador",
             "Se asigna el perfil de acceso (administrador, cajero o padre) que define las funcionalidades habilitadas."),
            ("CU-59", "Restringir acceso por perfil", "Sistema",
             "El sistema verifica en cada peticion que el usuario posee los permisos requeridos."),
            ("CU-60", "Activar autenticacion de dos factores (2FA)", "Administrador, Cajero",
             "El usuario habilita la verificacion adicional via TOTP para reforzar la seguridad de su cuenta."),
            ("CU-61", "Registrar accion critica en bitacora", "Sistema",
             "Cada operacion sensible queda vinculada al usuario autenticado con marca de tiempo."),
            ("CU-62", "Modificar o dar de baja usuario", "Administrador",
             "El administrador actualiza datos de acceso o desactiva la cuenta sin eliminar el historial."),
        ],
    },
]

ER_DIAGRAMS = [
    ("ER-01", "Modulo Usuarios y Autenticacion", "diag_17.png",
     "Usuarios del sistema con autenticacion, roles y permisos RBAC, sesiones activas, intentos "
     "de login, recuperacion de contrasena y soporte para autenticacion de dos factores (2FA)."),
    ("ER-02", "Modulo Clientes e Hijos", "diag_18.png",
     "Clientes (padres/responsables) con sus hijos/alumnos, cuenta corriente prepaga, historial "
     "de grados escolares y restricciones alimenticias por hijo."),
    ("ER-03", "Modulo Productos y Precios", "diag_19.png",
     "Catalogo de productos con categorias jerarquicas, multiples listas de precios y asignacion "
     "de impuestos por producto. Historico de precios para analisis de rentabilidad."),
    ("ER-04", "Modulo Core / Tarjetas y Saldo", "diag_20.png",
     "Tarjetas monedero vinculadas a hijos, cargas de saldo, consumos, medios de pago, limites de "
     "transaccion por rol y registro de alertas automaticas de saldo bajo."),
    ("ER-05", "Modulo Ventas / Punto de Venta", "diag_21.png",
     "Ventas con desglose de IVA, pagos multiples por venta, notas de credito y aplicacion de "
     "pagos a multiples ventas. Vinculacion con medios de pago y tarjetas."),
    ("ER-06", "Modulo Almuerzos", "diag_22.png",
     "Planes de almuerzo, suscripciones de hijos, registro de consumo diario y cuentas mensuales. "
     "Menu diario con platos disponibles y tipos de plan por periodo."),
    ("ER-07", "Modulo Inventario", "diag_23.png",
     "Stock por producto, ajustes con aprobacion de doble nivel, lotes con fecha de vencimiento, "
     "costos historicos y alertas automaticas de stock critico."),
    ("ER-08", "Modulo Compras y Proveedores", "diag_24.png",
     "Compras a proveedores con detalle por producto, pagos con aplicacion a multiples compras, "
     "notas de credito y cuenta corriente de proveedor completamente aislada de clientes."),
    ("ER-09", "Modulo Contabilidad / Caja / Facturacion", "diag_25.png",
     "Gestion de cajas fisicas, cierres de caja diarios con movimientos clasificados por medio "
     "de pago, facturas preimpresas legales y datos de la empresa emisora."),
    ("ER-10/11", "Modulos Notificaciones e Integraciones API", "diag_26.png",
     "Notificaciones de sistema y por email con plantillas parametrizadas, y gestion de "
     "proveedores de API externos con credenciales, webhooks y logs de auditoria."),
]

SEQ_DIAGRAMS = [
    ("SEQ-01", "Login y Autenticacion 2FA", "diag_27.png",
     "Flujo de autenticacion con JWT (SimpleJWT). Si el usuario tiene 2FA habilitado, debe "
     "ingresar el codigo TOTP antes de recibir el token de acceso."),
    ("SEQ-02", "Carga de Saldo en Tarjeta", "diag_28.png",
     "El cajero selecciona tarjeta y monto, el backend valida limites, crea la carga y actualiza "
     "el saldo disponible en tiempo real."),
    ("SEQ-03", "Registro de Venta con Tarjeta", "diag_29.png",
     "Registro de una venta pagada con tarjeta monedero. Incluye validacion de stock, calculo "
     "de IVA, descuento de saldo en tarjeta y registro del movimiento de caja."),
    ("SEQ-04", "Registro de Consumo de Almuerzo", "diag_30.png",
     "El cajero registra el consumo de almuerzo de un alumno. El sistema verifica la suscripcion "
     "vigente, registra el consumo y genera la cuenta mensual correspondiente."),
    ("SEQ-05", "Registro de Compra y Actualizacion de Stock", "diag_31.png",
     "El administrador registra una compra a proveedor. El backend actualiza automaticamente "
     "el stock de cada producto y registra la deuda en la cuenta corriente del proveedor."),
    ("SEQ-06", "Ajuste de Inventario (Solicitar y Aprobar)", "diag_32.png",
     "Flujo de doble paso: el usuario solicita un ajuste de inventario y el administrador lo "
     "aprueba o rechaza. Solo al aprobar se actualiza el stock fisico."),
    ("SEQ-07", "Emision de Factura Legal", "diag_33.png",
     "Flujo de emision de factura preimpresa para una carga de saldo o pago de almuerzo. "
     "El sistema calcula IVA automaticamente y valida la correlacion numerica del talonario."),
    ("SEQ-08", "Cierre de Caja", "diag_34.png",
     "Al final del dia, el cajero cuenta el efectivo fisico y cierra la caja. El sistema calcula "
     "la diferencia entre el monto esperado y el contado, generando el reporte de auditoria."),
]


# ════════════════════════════════════════════════════════════════════════════
#   CONSTRUCCION
# ════════════════════════════════════════════════════════════════════════════
def build_portada(doc):
    doc.add_paragraph().paragraph_format.space_before = Pt(30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sistema de Gestion — Cantina Tita")
    r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Documentacion UML Completa")
    r2.font.size = Pt(16); r2.font.bold = True; r2.font.color.rgb = GREEN

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        "Casos de Uso  |  Especificacion de Actores  |  Diagramas ER por Modulo  "
        "|  Secuencias  |  Arquitectura"
    )
    r3.font.size = Pt(9.5); r3.font.color.rgb = GRAY

    doc.add_paragraph()
    meta = [
        ("Documento",  "Documentacion UML — Especificacion Completa"),
        ("Version",    "1.0"),
        ("Fecha",      "Mayo 2026"),
        ("Proyecto",   "Sistema de Gestion Cantina Tita"),
        ("Modulos",    "10 modulos funcionales — 62 casos de uso"),
        ("Diagramas",  "10 ER + 10 CU + 8 Secuencia + Despliegue + Requerimientos"),
        ("Autores",    "Cantina Tita Development Team"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.style = "Table Grid"
    for i, (lbl, val) in enumerate(meta):
        row = tbl.rows[i]
        row.cells[0].text = lbl; row.cells[1].text = val
        cell_bg(row.cells[0], "1E3A5F"); cell_bg(row.cells[1], "E8F5E9")
        for r in row.cells[0].paragraphs[0].runs:
            r.font.color.rgb = WHITE; r.font.bold = True; r.font.size = Pt(9.5)
        for r in row.cells[1].paragraphs[0].runs:
            r.font.size = Pt(9.5)
    doc.add_paragraph()
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run("CONFIDENCIAL — USO INTERNO")
    r5.font.bold = True; r5.font.size = Pt(10); r5.font.color.rgb = NAVY
    doc.add_page_break()


def build_indice(doc):
    h1(doc, "Tabla de Contenidos")
    secciones = [
        ("Parte I",   "Diagramas de Casos de Uso"),
        ("  1.1",     "Introduccion"),
        ("  1.2",     "Especificacion de Actores"),
        ("  1.3",     "Diagramas de Casos de Uso (graficos por modulo)"),
        ("  1.4",     "Especificacion detallada por modulo (objetivos + CU + ER)"),
        ("Parte II",  "Diagramas Entidad-Relacion por Modulo"),
        ("  ER-01",   "Modulo Usuarios y Autenticacion"),
        ("  ER-02",   "Modulo Clientes e Hijos"),
        ("  ER-03",   "Modulo Productos y Precios"),
        ("  ER-04",   "Modulo Core / Tarjetas y Saldo"),
        ("  ER-05",   "Modulo Ventas / Punto de Venta"),
        ("  ER-06",   "Modulo Almuerzos"),
        ("  ER-07",   "Modulo Inventario"),
        ("  ER-08",   "Modulo Compras y Proveedores"),
        ("  ER-09",   "Modulo Contabilidad / Caja / Facturacion"),
        ("  ER-10/11","Modulo Notificaciones e Integraciones API"),
        ("Parte III", "Diagramas de Secuencia — Flujos Clave"),
        ("  SEQ-01",  "Login y Autenticacion 2FA"),
        ("  SEQ-02",  "Carga de Saldo en Tarjeta"),
        ("  SEQ-03",  "Registro de Venta con Tarjeta"),
        ("  SEQ-04",  "Registro de Consumo de Almuerzo"),
        ("  SEQ-05",  "Registro de Compra y Actualizacion de Stock"),
        ("  SEQ-06",  "Ajuste de Inventario (Solicitar y Aprobar)"),
        ("  SEQ-07",  "Emision de Factura Legal"),
        ("  SEQ-08",  "Cierre de Caja"),
        ("Parte IV",  "Diagramas de Arquitectura"),
        ("  IV.1",    "Diagrama de Despliegue"),
        ("  IV.2",    "Diagrama de Requerimientos"),
    ]
    for sec, nombre in secciones:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
        is_main = not sec.startswith("  ")
        r1 = p.add_run(f"{sec}  ")
        r1.font.bold = is_main; r1.font.size = Pt(10.5 if is_main else 10)
        if is_main: r1.font.color.rgb = NAVY
        r2 = p.add_run(nombre)
        r2.font.bold = is_main; r2.font.size = Pt(10.5 if is_main else 10)
        if is_main: r2.font.color.rgb = NAVY
    doc.add_page_break()


def build_parte_i(doc):
    section_divider(doc, "Parte I", "Diagramas de Casos de Uso")

    # 1.1 Introduccion
    h2(doc, "1.1 Introduccion")
    quote(doc,
        "El diagrama de casos de uso representa las interacciones entre los actores del "
        "sistema y las funcionalidades que este pone a disposicion de cada perfil. Cada "
        "actor agrupa a un conjunto de usuarios que comparten necesidades y niveles de "
        "acceso especificos, permitiendo una segmentacion clara de las operaciones criticas "
        "del negocio."
    )
    body(doc,
        "Los casos de uso se agruparon por modulo, conservando la correspondencia uno a uno "
        "con los objetivos especificos del trabajo. Para cada modulo se describe el objetivo "
        "funcional asociado, los actores involucrados, la lista de casos de uso y el diagrama "
        "entidad-relacion que modela los datos subyacentes."
    )

    # 1.2 Actores
    doc.add_paragraph()
    h2(doc, "1.2 Especificacion de Actores")
    body(doc,
        "El sistema Cantina Tita define cuatro actores principales. Cada actor representa "
        "un perfil con necesidades y niveles de acceso diferenciados. La jerarquia de roles "
        "garantiza que cada operacion sea ejecutada exclusivamente por el perfil autorizado."
    )
    doc.add_paragraph()
    table_actor(doc, ACTORES)
    body(doc,
        "El actor Sistema actua de forma automatizada como respuesta a eventos del negocio: "
        "cambios de stock, umbrales de saldo, cierres de caja y acciones de auditoria. "
        "No requiere intervencion humana directa y es responsable de mantener la consistencia "
        "e integridad de los datos en todo momento."
    )

    # 1.3 Diagramas graficos CU
    doc.add_page_break()
    h2(doc, "1.3 Diagramas de Casos de Uso por Modulo (graficos)")
    body(doc,
        "Los siguientes diagramas muestran visualmente las relaciones entre actores y "
        "funcionalidades para cada uno de los 11 modulos de la aplicacion Django."
    )
    cu_pages = [
        ("CU-01", "Usuarios y Autenticacion",       "diag_05.png"),
        ("CU-02", "Clientes e Hijos",               "diag_06.png"),
        ("CU-03", "Productos y Precios",             "diag_07.png"),
        ("CU-04", "Tarjetas y Saldo (Core)",         "diag_08.png"),
        ("CU-05", "Ventas / Punto de Venta",         "diag_09.png"),
        ("CU-06", "Almuerzos",                       "diag_10.png"),
        ("CU-07", "Inventario",                      "diag_11.png"),
        ("CU-08", "Compras",                         "diag_12.png"),
        ("CU-09", "Contabilidad y Caja",             "diag_13.png"),
        ("CU-10", "Notificaciones",                  "diag_14.png"),
        ("CU-11", "Integraciones API",               "diag_15.png"),
    ]
    for cuid, nombre, png in cu_pages:
        doc.add_paragraph()
        h3(doc, f"{cuid} — {nombre}")
        img(doc, png, width=6.0, caption=f"Diagrama de Casos de Uso — {nombre}")

    # 1.4 Especificacion por modulo
    doc.add_page_break()
    h2(doc, "1.4 Especificacion Detallada por Modulo")
    body(doc,
        "Para cada modulo se presenta: el objetivo especifico del trabajo, los actores "
        "involucrados, la tabla de casos de uso identificados y el diagrama entidad-relacion "
        "que modela los datos que soportan las funcionalidades descritas."
    )

    for mod in MODULOS:
        doc.add_page_break()
        h3(doc, f"Modulo {mod['num']}: {mod['titulo']}")

        label_val(doc, "Objetivo especifico", mod["objetivo"])
        label_val(doc, "Actores involucrados", mod["actores"])
        doc.add_paragraph()

        p = doc.add_paragraph()
        r = p.add_run("Casos de uso identificados:")
        r.font.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = NAVY
        doc.add_paragraph()

        table_cu(doc, mod["casos"])

        # Diagrama CU grafico del modulo mas relevante
        doc.add_paragraph()
        p2 = doc.add_paragraph()
        r2 = p2.add_run("Diagrama de casos de uso:")
        r2.font.bold = True; r2.font.size = Pt(10.5); r2.font.color.rgb = NAVY

        img(doc, mod["cu_diag"], width=5.8,
            caption=f"Diagrama CU — {mod['titulo']}")

        # Diagrama ER del modulo
        doc.add_paragraph()
        p3 = doc.add_paragraph()
        r3 = p3.add_run(f"Diagrama Entidad-Relacion — {mod['er_label']}:")
        r3.font.bold = True; r3.font.size = Pt(10.5); r3.font.color.rgb = NAVY

        p4 = doc.add_paragraph(mod["er_desc"])
        p4.paragraph_format.left_indent = Inches(0.2)
        p4.paragraph_format.space_after = Pt(6)
        for r in p4.runs:
            r.font.size = Pt(10); r.font.color.rgb = GRAY
        border_left(p4)

        img(doc, mod["er_diag"], width=6.0,
            caption=f"ER — {mod['er_label']}")


def build_parte_ii(doc):
    section_divider(doc, "Parte II", "Diagramas Entidad-Relacion por Modulo")

    body(doc,
        "Los diagramas entidad-relacion presentan la estructura de datos persistida en "
        "PostgreSQL para cada modulo del sistema. Las entidades representan tablas de la base "
        "de datos Django y las relaciones modelan claves foraneas, cardinalidades y "
        "restricciones de integridad referencial."
    )
    doc.add_paragraph()

    for er_id, nombre, png, desc in ER_DIAGRAMS:
        h2(doc, f"{er_id} — {nombre}")
        p = doc.add_paragraph(desc)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(8)
        for r in p.runs:
            r.font.size = Pt(10.5)
        border_left(p)
        img(doc, png, width=6.2, caption=f"Diagrama ER — {nombre}")
        doc.add_page_break()


def build_parte_iii(doc):
    section_divider(doc, "Parte III", "Diagramas de Secuencia — Flujos Clave")

    body(doc,
        "Los diagramas de secuencia representan los flujos de mensajes entre los componentes "
        "del sistema (frontend React, API Django, base de datos PostgreSQL y servicios externos) "
        "para los procesos de negocio mas criticos. Cada diagrama muestra el orden cronologico "
        "de las interacciones y las validaciones del sistema."
    )
    doc.add_paragraph()

    for seq_id, nombre, png, desc in SEQ_DIAGRAMS:
        h2(doc, f"{seq_id} — {nombre}")
        p = doc.add_paragraph(desc)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(8)
        for r in p.runs:
            r.font.size = Pt(10.5)
        border_left(p)
        img(doc, png, width=6.2, caption=f"Diagrama de Secuencia — {nombre}")
        doc.add_page_break()


def build_parte_iv(doc):
    section_divider(doc, "Parte IV", "Arquitectura del Sistema")

    h2(doc, "IV.1 Diagrama de Despliegue")
    body(doc,
        "Infraestructura de produccion: contenedores Docker sobre servidor Linux. El frontend "
        "(Nginx + React) recibe peticiones del navegador y del portal movil, el backend Django "
        "corre con Gunicorn como servidor WSGI, y los workers Celery procesan tareas asincronas "
        "(emails, alertas). Redis actua como broker de mensajes y cache. PostgreSQL es la base "
        "de datos central compartida por todos los modulos."
    )
    img(doc, "diag_03.png", width=6.2, caption="Diagrama de Despliegue — Infraestructura Docker")

    doc.add_page_break()
    h2(doc, "IV.2 Diagrama de Requerimientos")
    body(doc,
        "Arbol de requerimientos funcionales agrupados por dominio de negocio. Cada rama del "
        "mindmap corresponde a un modulo del sistema y lista las capacidades funcionales que "
        "ese modulo provee. Los requerimientos no funcionales (seguridad, performance, "
        "usabilidad) son transversales a todos los modulos."
    )
    img(doc, "diag_04.png", width=6.2, caption="Diagrama de Requerimientos — Mindmap funcional")


def build_resumen(doc):
    doc.add_page_break()
    h1(doc, "Resumen de Cobertura")
    body(doc, "Tabla consolidada de modulos, actores y cantidad de casos de uso especificados.")
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    set_repeat_header(hdr)
    for i, txt in enumerate(["Modulo", "Nombre", "Actores", "Casos de uso"]):
        c = hdr.cells[i]; c.text = txt
        cell_bg(c, "1E3A5F")
        for r in c.paragraphs[0].runs:
            r.font.color.rgb = WHITE; r.font.bold = True; r.font.size = Pt(9)

    for ri, mod in enumerate(MODULOS):
        row = tbl.add_row().cells
        row[0].text = f"M-{mod['num']:0>2}"
        row[1].text = mod["titulo"]
        row[2].text = mod["actores"]
        row[3].text = str(len(mod["casos"]))
        bg = "E8F5E9" if ri % 2 == 0 else "FFFFFF"
        for ci in range(4):
            cell_bg(row[ci], bg)
            for r in row[ci].paragraphs[0].runs: r.font.size = Pt(9)
        for r in row[0].paragraphs[0].runs:
            r.font.bold = True; r.font.color.rgb = NAVY
        for r in row[1].paragraphs[0].runs: r.font.bold = True
        for r in row[3].paragraphs[0].runs:
            r.font.bold = True; r.font.color.rgb = GREEN

    doc.add_paragraph()
    total = sum(len(m["casos"]) for m in MODULOS)
    p = doc.add_paragraph()
    r1 = p.add_run("Total de casos de uso: ")
    r1.font.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = NAVY
    r2 = p.add_run(str(total))
    r2.font.bold = True; r2.font.size = Pt(14); r2.font.color.rgb = GREEN
    r3 = p.add_run("   |   Actores definidos: ")
    r3.font.bold = True; r3.font.size = Pt(12); r3.font.color.rgb = NAVY
    r4 = p.add_run(str(len(ACTORES)))
    r4.font.bold = True; r4.font.size = Pt(14); r4.font.color.rgb = GREEN
    r5 = p.add_run("   |   Modulos cubiertos: ")
    r5.font.bold = True; r5.font.size = Pt(12); r5.font.color.rgb = NAVY
    r6 = p.add_run(str(len(MODULOS)))
    r6.font.bold = True; r6.font.size = Pt(14); r6.font.color.rgb = GREEN


# ── MAIN ──────────────────────────────────────────────────────────────────────
def main():
    doc = make_doc()
    print("Construyendo documento completo...")
    build_portada(doc)
    print("  [Portada]")
    build_indice(doc)
    print("  [Indice]")
    build_parte_i(doc)
    print("  [Parte I: Casos de Uso]")
    build_parte_ii(doc)
    print("  [Parte II: Diagramas ER]")
    build_parte_iii(doc)
    print("  [Parte III: Secuencias]")
    build_parte_iv(doc)
    print("  [Parte IV: Arquitectura]")
    build_resumen(doc)
    print("  [Resumen]")
    doc.save(str(OUT))
    total_cu = sum(len(m["casos"]) for m in MODULOS)
    print(f"\nListo: {OUT}")
    print(f"  Modulos: {len(MODULOS)} | CU: {total_cu} | ER: {len(ER_DIAGRAMS)} | SEQ: {len(SEQ_DIAGRAMS)}")


if __name__ == "__main__":
    main()
