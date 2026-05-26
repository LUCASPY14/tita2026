#!/usr/bin/env python3
"""
generar_docx.py - Convierte documentacion-uml.html -> documentacion-uml.docx
Renderiza cada diagrama Mermaid a PNG con mmdc y los embebe en el DOCX.
Uso: python generar_docx.py
"""

import subprocess
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x1E, 0x3A, 0x5F)
GREEN = RGBColor(0x27, 0xAE, 0x60)
GRAY  = RGBColor(0x88, 0x88, 0x88)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

DOCS_DIR = Path(r"d:\tita2026\docs\diagramas")
HTML_IN  = DOCS_DIR / "documentacion-uml.html"
DOCX_OUT = DOCS_DIR / "documentacion-uml.docx"
IMG_DIR  = DOCS_DIR / "diagrams_png"


# ── OOXML helpers ─────────────────────────────────────────────────────────────
def cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def left_border_green(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "18")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "27AE60")
    pBdr.append(left)
    pPr.append(pBdr)


def bottom_border_green(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "12")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "27AE60")
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Mermaid rendering ─────────────────────────────────────────────────────────
def render_mermaid(code: str, slug: str) -> Path | None:
    IMG_DIR.mkdir(exist_ok=True)
    out = IMG_DIR / f"{slug}.png"
    mmd = IMG_DIR / f"{slug}.mmd"
    mmd.write_text(code.strip(), encoding="utf-8")
    try:
        cmd = (
            f'npx @mermaid-js/mermaid-cli'
            f' -i "{mmd}" -o "{out}"'
            f' --backgroundColor white --width 1400 --scale 2'
        )
        res = subprocess.run(
            cmd,
            capture_output=True, text=True,
            cwd=str(DOCS_DIR), timeout=120,
            shell=True,
        )
        mmd.unlink(missing_ok=True)
        if out.exists() and out.stat().st_size > 0:
            print(f"    OK {slug}.png")
            return out
        print(f"    ERR {slug}: {res.stderr[:200]}")
        return None
    except Exception as e:
        mmd.unlink(missing_ok=True)
        print(f"    EXC {slug}: {e}")
        return None


# ── Document setup ────────────────────────────────────────────────────────────
def make_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.orientation   = WD_ORIENT.PORTRAIT
    sec.left_margin   = Inches(0.9)
    sec.right_margin  = Inches(0.9)
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10)
    return doc


def add_h(doc, text, level=2, color=NAVY, align=WD_ALIGN_PARAGRAPH.LEFT, border=False):
    p = doc.add_heading(text, level)
    p.alignment = align
    for r in p.runs:
        r.font.color.rgb = color
    if border:
        bottom_border_green(p)
    return p


def desc_para(doc, text: str):
    p = doc.add_paragraph(text.strip())
    p.paragraph_format.left_indent  = Inches(0.15)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY
    left_border_green(p)


def img_para(doc, img_path: Path, w=6.4):
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img_path), width=Inches(w))
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(8)
    except Exception as e:
        doc.add_paragraph(f"[Imagen no disponible: {e}]")


def code_block(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text[:3000])
    r.font.name = "Courier New"
    r.font.size = Pt(7)
    r.font.color.rgb = GRAY
    p.paragraph_format.left_indent = Inches(0.2)


def html_table_to_docx(doc, html_tbl):
    rows_h = html_tbl.find_all("tr")
    if not rows_h:
        return
    cols = max(len(r.find_all(["td", "th"])) for r in rows_h)
    tbl = doc.add_table(rows=0, cols=cols)
    tbl.style = "Table Grid"
    for ri, tr in enumerate(rows_h):
        cells_h = tr.find_all(["td", "th"])
        row = tbl.add_row()
        is_hdr = bool(cells_h) and cells_h[0].name == "th"
        for ci, ch in enumerate(cells_h):
            if ci >= cols:
                break
            cell = row.cells[ci]
            cell.text = ch.get_text(" ", strip=True)
            p = cell.paragraphs[0]
            if is_hdr:
                cell_bg(cell, "1E3A5F")
                for r in p.runs:
                    r.font.color.rgb = WHITE
                    r.font.bold = True
                    r.font.size = Pt(8)
            else:
                cell_bg(cell, "F8F9FA" if ri % 2 == 0 else "FFFFFF")
                for r in p.runs:
                    r.font.size = Pt(8)
    doc.add_paragraph()


# ── Cover ─────────────────────────────────────────────────────────────────────
def build_cover(doc):
    doc.add_paragraph().paragraph_format.space_before = Pt(36)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sistema de Gestion")
    r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Cantina Tita")
    r2.font.size = Pt(16); r2.font.bold = True; r2.font.color.rgb = GREEN

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Documentacion UML Completa")
    r3.font.size = Pt(14); r3.font.bold = True; r3.font.color.rgb = NAVY

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Casos de Uso  |  Entidad-Relacion  |  Secuencia  |  Requerimientos  |  Despliegue")
    r4.font.size = Pt(9); r4.font.color.rgb = GRAY

    doc.add_paragraph()

    meta = [
        ("Proyecto",  "Sistema de Gestion Cantina Tita"),
        ("Version",   "1.0.0"),
        ("Fecha",     "Mayo 2026"),
        ("Stack",     "Django 4.2 / React 19 / PostgreSQL / Docker"),
        ("Modulos",   "11 aplicaciones Django"),
        ("Autor",     "Cantina Tita Development Team"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.style = "Table Grid"
    for i, (lbl, val) in enumerate(meta):
        row = tbl.rows[i]
        row.cells[0].text = lbl
        row.cells[1].text = val
        cell_bg(row.cells[0], "1E3A5F")
        cell_bg(row.cells[1], "E8F5E9")
        for r in row.cells[0].paragraphs[0].runs:
            r.font.color.rgb = WHITE; r.font.bold = True; r.font.size = Pt(9)
        for r in row.cells[1].paragraphs[0].runs:
            r.font.size = Pt(9)

    doc.add_paragraph()
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run("CONFIDENCIAL - USO INTERNO")
    r5.font.bold = True; r5.font.size = Pt(10); r5.font.color.rgb = NAVY

    doc.add_page_break()


# ── TOC ───────────────────────────────────────────────────────────────────────
def build_toc(doc, toc_page):
    add_h(doc, "Tabla de Contenidos", level=1, color=NAVY)
    for child in toc_page.children:
        if not hasattr(child, "get"):
            continue
        classes = child.get("class", [])
        txt = child.get_text(strip=True)
        if not txt or "footer" in classes:
            continue
        if "toc-section" in classes:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            r = p.add_run(txt)
            r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY
        elif "toc-item" in classes:
            spans = child.find_all("span")
            if len(spans) >= 2:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_before = Pt(1)
                r1 = p.add_run(spans[0].get_text(strip=True))
                r1.font.size = Pt(9)
                r2 = p.add_run(f"  .........  {spans[1].get_text(strip=True)}")
                r2.font.size = Pt(9); r2.font.color.rgb = GREEN
    doc.add_page_break()


# ── Diagram page ──────────────────────────────────────────────────────────────
def build_diag_page(doc, page, idx: int):
    header  = page.find(class_="diag-header")
    desc_el = page.find(class_="diag-desc")
    merm_el = page.find("pre", class_="mermaid")
    spec_el = page.find("table", class_="spec-table")

    title = ""
    dtag  = ""
    if header:
        h2el = header.find("h2")
        title = h2el.get_text(strip=True) if h2el else ""
        dtag_el = header.find(class_="dtag")
        dtag = dtag_el.get_text(strip=True) if dtag_el else ""

    if title:
        add_h(doc, title, level=2, color=NAVY, border=True)
    if dtag:
        tp = doc.add_paragraph()
        r = tp.add_run(f"[ {dtag} ]")
        r.font.size = Pt(8); r.font.bold = True; r.font.color.rgb = GREEN

    if desc_el:
        desc_text = desc_el.get_text(" ", strip=True)
        if desc_text:
            desc_para(doc, desc_text)

    if merm_el:
        code = merm_el.get_text()
        slug = f"diag_{idx:02d}"
        img = render_mermaid(code, slug)
        if img:
            img_para(doc, img)
        else:
            code_block(doc, code)

    if spec_el:
        html_table_to_docx(doc, spec_el)

    doc.add_page_break()


# ── Summary page ──────────────────────────────────────────────────────────────
def build_summary(doc, page):
    add_h(doc, "Resumen de Documentacion UML", level=2, color=NAVY, border=True)
    spec_el = page.find("table", class_="spec-table")
    if spec_el:
        html_table_to_docx(doc, spec_el)
    for p_el in page.find_all("p"):
        txt = p_el.get_text(strip=True)
        if txt:
            doc.add_paragraph(txt)


# ── MAIN ──────────────────────────────────────────────────────────────────────
def main():
    print("Leyendo HTML...")
    html = HTML_IN.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "lxml")
    pages = soup.find_all("div", class_="page")
    total = len(pages)
    print(f"Paginas encontradas: {total}")

    doc = make_doc()

    print("[1] Portada")
    build_cover(doc)

    print("[2] Tabla de contenidos")
    build_toc(doc, pages[1])

    # pages[2] .. pages[-2] son diagrams; pages[-1] es resumen
    for i, page in enumerate(pages[2:-1], start=3):
        h2el = page.find("h2")
        title = h2el.get_text(strip=True)[:55] if h2el else f"pag {i}"
        print(f"[{i:2d}] {title}")
        build_diag_page(doc, page, idx=i)

    print(f"[{total}] Resumen")
    build_summary(doc, pages[-1])

    print(f"\nGuardando -> {DOCX_OUT}")
    doc.save(str(DOCX_OUT))
    print(f"Listo: {DOCX_OUT}")


if __name__ == "__main__":
    main()
