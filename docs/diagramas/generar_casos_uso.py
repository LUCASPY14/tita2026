#!/usr/bin/env python3
"""
generar_casos_uso.py
Genera casos_de_uso.docx: texto completo de Diagramas de Casos de Uso
para el Sistema de Gestion Cantina Tita.
Uso: python generar_casos_uso.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
GREEN  = RGBColor(0x27, 0xAE, 0x60)
DARK   = RGBColor(0x2C, 0x3E, 0x50)
GRAY   = RGBColor(0x66, 0x66, 0x66)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xE8, 0xF5, 0xE9)

OUT = Path(r"d:\tita2026\docs\diagramas\casos_de_uso.docx")


# ── OOXML helpers ──────────────────────────────────────────────────────────────
def cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def bottom_border(para, color="27AE60", sz=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def left_border(para, color="27AE60", sz=18):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(sz))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


# ── Document factory ───────────────────────────────────────────────────────────
def make_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.orientation   = WD_ORIENT.PORTRAIT
    sec.left_margin   = Inches(1.0)
    sec.right_margin  = Inches(1.0)
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    return doc


# ── Paragraph helpers ──────────────────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(16)
    bottom_border(p, "1E3A5F", 16)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(13)
    bottom_border(p, "27AE60", 10)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = DARK
        r.font.size = Pt(11)
    return p


def body(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def quote_block(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    for r in p.runs:
        r.font.size = Pt(10.5)
        r.font.italic = True
        r.font.color.rgb = DARK
    left_border(p, "27AE60", 18)


def bullet(doc, text, indent=0.3):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def label_value(doc, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.font.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(value)
    r2.font.size = Pt(10.5)


# ── Tables ─────────────────────────────────────────────────────────────────────
def actor_table(doc, actors):
    """
    actors = list of (id, nombre, descripcion, permisos)
    """
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(["ID", "Actor", "Descripcion", "Nivel de acceso"]):
        hdr[i].text = txt
        cell_bg(hdr[i], "1E3A5F")
        for r in hdr[i].paragraphs[0].runs:
            r.font.color.rgb = WHITE
            r.font.bold = True
            r.font.size = Pt(9)

    for ri, (aid, nombre, desc, acceso) in enumerate(actors):
        row = tbl.add_row().cells
        row[0].text = aid
        row[1].text = nombre
        row[2].text = desc
        row[3].text = acceso
        bg = "E8F5E9" if ri % 2 == 0 else "FFFFFF"
        for ci in range(4):
            cell_bg(row[ci], bg)
            for r in row[ci].paragraphs[0].runs:
                r.font.size = Pt(9.5)
        # ID en negrita
        for r in row[0].paragraphs[0].runs:
            r.font.bold = True
            r.font.color.rgb = NAVY
        # nombre en negrita
        for r in row[1].paragraphs[0].runs:
            r.font.bold = True

    doc.add_paragraph()


def cu_table(doc, casos):
    """
    casos = list of (id, nombre, actores, descripcion)
    """
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(["ID", "Caso de Uso", "Actor(es)", "Descripcion"]):
        hdr[i].text = txt
        cell_bg(hdr[i], "1E3A5F")
        for r in hdr[i].paragraphs[0].runs:
            r.font.color.rgb = WHITE
            r.font.bold = True
            r.font.size = Pt(9)

    for ri, (cid, nombre, actores, desc) in enumerate(casos):
        row = tbl.add_row().cells
        row[0].text = cid
        row[1].text = nombre
        row[2].text = actores
        row[3].text = desc
        bg = "F0F7F4" if ri % 2 == 0 else "FFFFFF"
        for ci in range(4):
            cell_bg(row[ci], bg)
            for r in row[ci].paragraphs[0].runs:
                r.font.size = Pt(9)
        for r in row[0].paragraphs[0].runs:
            r.font.bold = True
            r.font.color.rgb = GREEN
        for r in row[1].paragraphs[0].runs:
            r.font.bold = True

    doc.add_paragraph()


# ══ CONTENIDO ══════════════════════════════════════════════════════════════════

INTRO = (
    "El diagrama de casos de uso representa las interacciones entre los actores del "
    "sistema y las funcionalidades que este pone a disposicion de cada perfil. Cada "
    "actor agrupa a un conjunto de usuarios que comparten necesidades y niveles de "
    "acceso especificos, permitiendo una segmentacion clara de las operaciones criticas "
    "del negocio."
)

ACTORES = [
    ("A-01", "Administrador",
     "Usuario con acceso total al sistema. Gestiona configuracion general, usuarios, "
     "reportes avanzados, facturacion y parametros del negocio.",
     "Acceso completo (CRUD en todos los modulos)"),
    ("A-02", "Cajero",
     "Operador del punto de venta. Registra ventas de cantina y almuerzo, procesa "
     "pagos, abre/cierra caja y gestiona tarjetas prepagas en el POS.",
     "Acceso operativo: POS, caja, tarjetas, comprobantes"),
    ("A-03", "Padre / Responsable",
     "Cliente externo que accede al portal para consultar el saldo prepago de su hijo, "
     "configurar topes de gasto diarios y recibir notificaciones automaticas.",
     "Acceso restringido: portal web, consulta y configuracion propia"),
    ("A-04", "Sistema",
     "Actor secundario automatizado. Ejecuta triggers de base de datos, actualiza stock, "
     "genera alertas de saldo bajo, registra bitacoras de auditoria y calcula IVA.",
     "Acciones internas automatizadas, sin intervencion humana"),
]

MODULOS = [
    {
        "num": "1",
        "titulo": "Gestion de Clientes",
        "objetivo": (
            "Disenar e implementar un modulo de gestion de clientes que incorpore la "
            "vinculacion directa cliente-hijo y el control de cuentas corrientes prepagas, "
            "permitiendo a los padres configurar topes de gasto diarios y recibir alertas "
            "automaticas de saldo bajo."
        ),
        "actores": "Administrador, Padre / Responsable, Sistema",
        "casos": [
            ("CU-01", "Registrar cliente",
             "Administrador",
             "El administrador ingresa los datos del responsable (nombre, RUC/CI, contacto) "
             "y crea el perfil en el sistema."),
            ("CU-02", "Vincular hijo a cliente",
             "Administrador",
             "Se asocia uno o mas hijos (estudiantes) al cliente responsable, estableciendo "
             "la relacion de dependencia para el control de consumos."),
            ("CU-03", "Configurar tope de gasto diario",
             "Padre / Responsable",
             "El padre define el limite maximo de gasto diario que puede efectuar su hijo "
             "en el punto de venta."),
            ("CU-04", "Consultar cuenta corriente",
             "Padre / Responsable, Administrador",
             "Se visualiza el historial de recargas, consumos y saldo disponible de la "
             "cuenta prepaga del hijo."),
            ("CU-05", "Cargar saldo prepago",
             "Administrador, Cajero",
             "El operador acredita un monto en la cuenta corriente del hijo mediante "
             "efectivo u otro medio de pago registrado."),
            ("CU-06", "Recibir alerta de saldo bajo",
             "Sistema, Padre / Responsable",
             "El sistema detecta que el saldo disponible cae por debajo del umbral "
             "configurado y envia notificacion automatica al padre."),
            ("CU-07", "Consultar historial de consumos",
             "Padre / Responsable, Administrador",
             "Lista cronologica de todas las transacciones realizadas por el hijo en "
             "el punto de venta."),
        ],
    },
    {
        "num": "2",
        "titulo": "Administracion de Proveedores",
        "objetivo": (
            "Desarrollar un modulo de administracion de proveedores que gestione de manera "
            "totalmente aislada e independiente sus respectivas cuentas corrientes (cuentas "
            "por pagar), garantizando la separacion estricta con respecto a los saldos de "
            "los clientes."
        ),
        "actores": "Administrador",
        "casos": [
            ("CU-08", "Registrar proveedor",
             "Administrador",
             "El administrador crea el perfil del proveedor con razon social, RUC, "
             "contacto y condiciones de pago."),
            ("CU-09", "Consultar cuenta corriente de proveedor",
             "Administrador",
             "Se visualiza el saldo de la deuda pendiente con el proveedor, separado "
             "completamente de las cuentas de clientes."),
            ("CU-10", "Registrar deuda por compra",
             "Administrador, Sistema",
             "Al procesar una compra, el sistema acredita automaticamente la deuda "
             "correspondiente en la cuenta del proveedor."),
            ("CU-11", "Registrar pago a proveedor",
             "Administrador",
             "El administrador registra el pago total o parcial de una deuda, actualizando "
             "el saldo de la cuenta por pagar."),
            ("CU-12", "Consultar historial de transacciones de proveedor",
             "Administrador",
             "Listado de todas las compras y pagos asociados a un proveedor especifico, "
             "con fechas e importes."),
        ],
    },
    {
        "num": "3",
        "titulo": "Arqueo y Cierre de Caja Diario",
        "objetivo": (
            "Disenar e implementar un modulo de arqueo y cierre de caja diario que registre "
            "de forma automatizada los ingresos por multiples medios de pago, permitiendo "
            "conciliar el efectivo fisico con las transacciones del sistema y documentar "
            "cualquier desvio financiero para su posterior auditoria."
        ),
        "actores": "Cajero, Administrador, Sistema",
        "casos": [
            ("CU-13", "Abrir caja del dia",
             "Cajero, Administrador",
             "El operador registra el monto inicial en efectivo y habilita la caja para "
             "recibir transacciones."),
            ("CU-14", "Registrar ingreso por medio de pago",
             "Sistema",
             "Cada venta acredita automaticamente el monto en el registro de la caja "
             "segun el medio de pago utilizado (efectivo, tarjeta, QR)."),
            ("CU-15", "Registrar egreso / retiro",
             "Cajero, Administrador",
             "Se documenta cualquier salida de dinero de la caja (vuelto, retiro "
             "autorizado) con su justificacion."),
            ("CU-16", "Realizar arqueo fisico",
             "Cajero",
             "El cajero ingresa el total contado fisicamente en caja al momento del cierre."),
            ("CU-17", "Cerrar caja y calcular diferencia",
             "Cajero, Sistema",
             "El sistema calcula la diferencia entre el monto esperado (segun registros) "
             "y el contado fisico, generando el reporte de desvio."),
            ("CU-18", "Generar reporte de cierre de caja",
             "Administrador, Sistema",
             "El sistema produce el comprobante del cierre con detalle por medio de pago, "
             "ingresos, egresos y diferencia, disponible para auditoria."),
        ],
    },
    {
        "num": "4",
        "titulo": "Punto de Venta (POS)",
        "objetivo": (
            "Disenar una interfaz de Punto de Venta (POS) optimizada para la alta "
            "concurrencia de los recreos, que diferencie de manera agil las operaciones "
            "de 'almuerzo' y 'cantina/recreo', con el fin de minimizar el tiempo medio "
            "de atencion por estudiante."
        ),
        "actores": "Cajero, Sistema",
        "casos": [
            ("CU-19", "Registrar venta de cantina / recreo",
             "Cajero",
             "El cajero selecciona productos del catalogo y procesa la venta en modalidad "
             "cantina, aplicando precio de lista correspondiente."),
            ("CU-20", "Registrar consumo de almuerzo",
             "Cajero",
             "El sistema registra el consumo del plan de almuerzo del estudiante y debita "
             "el importe de su cuenta prepaga o plan semanal."),
            ("CU-21", "Procesar pago con tarjeta prepaga",
             "Cajero, Sistema",
             "El sistema debita el monto de venta del saldo disponible en la tarjeta del "
             "estudiante, verificando topes y disponibilidad."),
            ("CU-22", "Procesar pago en efectivo u otro medio",
             "Cajero",
             "El cajero registra el pago con efectivo, tarjeta de debito/credito u otro "
             "medio habilitado, registrando el vuelto si corresponde."),
            ("CU-23", "Anular venta",
             "Cajero, Administrador",
             "Se revierte una venta registrada por error, devolviendo el saldo al cliente "
             "y ajustando el stock de forma automatica."),
            ("CU-24", "Emitir comprobante de venta",
             "Sistema",
             "El sistema genera el ticket o comprobante de la transaccion para entrega al "
             "cliente."),
            ("CU-25", "Consultar historial de ventas del turno",
             "Cajero, Administrador",
             "Listado de todas las transacciones procesadas durante el turno activo de la "
             "caja."),
        ],
    },
    {
        "num": "5",
        "titulo": "Compras",
        "objetivo": (
            "Construir un modulo de compras articulado con el inventario centralizado y el "
            "registro de proveedores, que permita procesar comprobantes de recepcion de "
            "insumos y actualizar de forma automatica los costos operativos de los "
            "productos."
        ),
        "actores": "Administrador, Sistema",
        "casos": [
            ("CU-26", "Crear orden de compra",
             "Administrador",
             "El administrador genera una solicitud formal de compra de insumos a un "
             "proveedor, especificando productos y cantidades."),
            ("CU-27", "Recepcionar insumos",
             "Administrador",
             "Se registra la recepcion fisica de los insumos segun el comprobante del "
             "proveedor (factura o remision)."),
            ("CU-28", "Actualizar stock automaticamente",
             "Sistema",
             "Al confirmar la recepcion, el sistema incrementa las unidades disponibles "
             "en el inventario centralizado."),
            ("CU-29", "Actualizar costo operativo del producto",
             "Sistema",
             "El sistema recalcula el costo de los productos afectados en funcion del "
             "precio unitario de la compra recepcionada."),
            ("CU-30", "Vincular compra a cuenta de proveedor",
             "Sistema",
             "La compra queda asociada al proveedor, generando automaticamente la deuda "
             "en su cuenta corriente."),
            ("CU-31", "Consultar historial de compras",
             "Administrador",
             "Listado de todas las compras realizadas, filtrable por proveedor, fecha "
             "o producto."),
        ],
    },
    {
        "num": "6",
        "titulo": "Facturacion Legal",
        "objetivo": (
            "Incorporar un motor de facturacion legal que aplique las normativas fiscales "
            "vigentes, validando el campo estructurado RUC_CI, automatizando el calculo "
            "de las tasas del IVA y garantizando la correlacion numerica exacta de las "
            "facturas preimpresas."
        ),
        "actores": "Administrador, Cajero, Sistema",
        "casos": [
            ("CU-32", "Emitir factura legal",
             "Administrador, Cajero",
             "Se registra una factura preimpresa vinculada a una venta o carga de saldo, "
             "ingresando el numero de comprobante y los datos del cliente."),
            ("CU-33", "Validar RUC / CI del cliente",
             "Sistema",
             "El sistema verifica el formato y la estructura del RUC o CI del cliente "
             "antes de emitir la factura, rechazando datos invalidos."),
            ("CU-34", "Calcular IVA automaticamente",
             "Sistema",
             "El motor fiscal calcula IVA 10%, IVA 5% y monto exento segun la "
             "naturaleza del bien o servicio facturado."),
            ("CU-35", "Verificar correlacion numerica",
             "Sistema",
             "El sistema garantiza que el numero de factura registrado no este duplicado "
             "y respete la secuencia numerica de los talonarios."),
            ("CU-36", "Anular factura",
             "Administrador",
             "El administrador anula una factura emitida por error, registrando el motivo "
             "y bloqueando el numero para reutilizacion."),
            ("CU-37", "Consultar facturas emitidas",
             "Administrador",
             "Listado de todas las facturas del periodo, filtrable por numero, cliente "
             "o fecha."),
        ],
    },
    {
        "num": "7",
        "titulo": "Tarjetas y Saldo Prepago",
        "objetivo": (
            "Desarrollar un mecanismo de control para tarjetas de uso exclusivo de la "
            "cantina, que habilite el procesamiento rapido de recargas de saldo y el debito "
            "automatizado de transacciones en el punto de venta."
        ),
        "actores": "Administrador, Cajero, Padre / Responsable, Sistema",
        "casos": [
            ("CU-38", "Activar tarjeta prepaga",
             "Administrador",
             "El administrador registra una nueva tarjeta en el sistema, la vincula al hijo "
             "correspondiente y la deja habilitada para operar."),
            ("CU-39", "Recargar saldo en tarjeta",
             "Cajero, Administrador",
             "El operador acredita un importe en la tarjeta del estudiante, registrando "
             "el medio de pago utilizado por el padre."),
            ("CU-40", "Debitar saldo en transaccion POS",
             "Sistema",
             "Al procesar una venta, el sistema descuenta automaticamente el importe del "
             "saldo disponible en la tarjeta, verificando disponibilidad y topes diarios."),
            ("CU-41", "Bloquear tarjeta",
             "Administrador, Padre / Responsable",
             "Se inhabilita temporalmente una tarjeta por extravío, robo o solicitud "
             "del responsable."),
            ("CU-42", "Consultar saldo y movimientos de tarjeta",
             "Padre / Responsable, Administrador",
             "Se visualiza el saldo actual y el historial de recargas y consumos de la "
             "tarjeta."),
            ("CU-43", "Generar alerta de saldo insuficiente",
             "Sistema",
             "El sistema detecta que el saldo de la tarjeta es menor al umbral configurado "
             "y notifica al padre."),
        ],
    },
    {
        "num": "8",
        "titulo": "Reportes Analiticos",
        "objetivo": (
            "Implementar un panel de reportes analiticos que consolide datos historicos de "
            "ventas e inventarios (filtrados por fecha, cajero, rotacion y rentabilidad), "
            "proveyendo a la administracion herramientas visuales para la toma de "
            "decisiones."
        ),
        "actores": "Administrador",
        "casos": [
            ("CU-44", "Generar reporte de ventas por periodo",
             "Administrador",
             "El sistema consolida todas las ventas del rango de fechas seleccionado, "
             "mostrando totales por dia, semana o mes."),
            ("CU-45", "Filtrar ventas por cajero",
             "Administrador",
             "Se desglosan las ventas segun el operador que las proceso, permitiendo "
             "evaluar el desempeno individual."),
            ("CU-46", "Consultar rotacion de productos",
             "Administrador",
             "El reporte muestra los productos mas y menos vendidos en el periodo, "
             "facilitando decisiones de compra y carta."),
            ("CU-47", "Analizar rentabilidad por producto",
             "Administrador",
             "El sistema calcula margen de ganancia unitario y total comparando precio "
             "de venta con costo de compra actualizado."),
            ("CU-48", "Visualizar dashboard de indicadores",
             "Administrador",
             "Panel de control con metricas clave: ventas del dia, saldo en caja, "
             "productos con stock critico y alertas activas."),
            ("CU-49", "Exportar reporte historico",
             "Administrador",
             "El administrador descarga el reporte seleccionado en formato Excel o PDF "
             "para uso externo o archivo."),
        ],
    },
    {
        "num": "9",
        "titulo": "Arquitectura de Base de Datos e Integridad",
        "objetivo": (
            "Disenar la arquitectura de la base de datos incorporando triggers y restricciones "
            "de integridad atomica para automatizar el control de stock en las ventas, "
            "asegurar la consistencia financiera en los cierres de caja y alimentar "
            "bitacoras de auditoria interna."
        ),
        "actores": "Sistema, Administrador",
        "casos": [
            ("CU-50", "Ejecutar trigger de control de stock en ventas",
             "Sistema",
             "Al registrar una venta, el trigger descuenta automaticamente las unidades "
             "del inventario y bloquea la operacion si el stock es insuficiente."),
            ("CU-51", "Validar integridad atomica en transacciones",
             "Sistema",
             "El motor de base de datos garantiza que cada operacion (venta, carga de "
             "saldo, compra) sea completamente exitosa o completamente revertida."),
            ("CU-52", "Actualizar cuentas corrientes en cierre de caja",
             "Sistema",
             "Los triggers consolidan los movimientos del dia en las cuentas corrientes "
             "de clientes y proveedores al cerrar la caja."),
            ("CU-53", "Registrar evento en bitacora de auditoria",
             "Sistema",
             "Cada accion critica (login, anulacion, cierre de caja, modificacion de precio) "
             "queda registrada en la bitacora con usuario, fecha y hora."),
            ("CU-54", "Consultar bitacora de auditoria",
             "Administrador",
             "El administrador accede al historial completo de acciones criticas para "
             "resolver discrepancias o investigar inconsistencias."),
            ("CU-55", "Gestionar restricciones de integridad referencial",
             "Sistema",
             "El sistema aplica claves foraneas y restricciones de unicidad para prevenir "
             "duplicados, datos huerfanos e inconsistencias entre modulos."),
        ],
    },
    {
        "num": "10",
        "titulo": "Control de Accesos y Usuarios",
        "objetivo": (
            "Implementar un modelo de control de accesos basado en roles jerarquicos "
            "(administrador, cajero, padre) que restrinja las funciones del sistema segun "
            "el perfil del usuario y registre las acciones criticas para la resolucion de "
            "discrepancias."
        ),
        "actores": "Administrador, Cajero, Padre / Responsable, Sistema",
        "casos": [
            ("CU-56", "Registrar usuario en el sistema",
             "Administrador",
             "El administrador crea la cuenta de un nuevo usuario asignando credenciales, "
             "rol y permisos correspondientes."),
            ("CU-57", "Autenticar usuario con JWT",
             "Sistema",
             "El sistema valida las credenciales del usuario y emite un token JWT firmado "
             "que habilita el acceso a las funciones del perfil."),
            ("CU-58", "Asignar rol jerarquico",
             "Administrador",
             "Se asigna el perfil de acceso (administrador, cajero o padre) que define las "
             "funcionalidades habilitadas para el usuario."),
            ("CU-59", "Restringir acceso por perfil",
             "Sistema",
             "El sistema verifica en cada peticion que el usuario posee los permisos "
             "requeridos, denegando el acceso si no corresponden a su rol."),
            ("CU-60", "Activar autenticacion de dos factores (2FA)",
             "Administrador, Cajero",
             "El usuario habilita la verificacion adicional via TOTP para reforzar la "
             "seguridad de su cuenta."),
            ("CU-61", "Registrar accion critica en bitacora",
             "Sistema",
             "Cada operacion sensible queda vinculada al usuario autenticado con marca "
             "de tiempo, para auditoria y resolucion de discrepancias."),
            ("CU-62", "Modificar o dar de baja usuario",
             "Administrador",
             "El administrador actualiza datos de acceso o desactiva la cuenta de un "
             "usuario sin eliminar su historial de operaciones."),
        ],
    },
]


# ══ CONSTRUCCION DEL DOCUMENTO ═════════════════════════════════════════════════
def build(doc: Document):

    # ── Portada ────────────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_before = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sistema de Gestion — Cantina Tita")
    r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Diagramas de Casos de Uso")
    r2.font.size = Pt(16); r2.font.bold = True; r2.font.color.rgb = GREEN

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Especificacion de Actores y Casos de Uso por Modulo")
    r3.font.size = Pt(11); r3.font.color.rgb = GRAY

    doc.add_paragraph()
    meta = [
        ("Documento",  "Especificacion de Casos de Uso"),
        ("Version",    "1.0"),
        ("Fecha",      "Mayo 2026"),
        ("Proyecto",   "Sistema de Gestion Cantina Tita"),
        ("Modulos",    "10 modulos funcionales — 62 casos de uso"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.style = "Table Grid"
    for i, (lbl, val) in enumerate(meta):
        row = tbl.rows[i]
        row.cells[0].text = lbl
        row.cells[1].text = val
        cell_bg(row.cells[0], "1E3A5F")
        cell_bg(row.cells[1], "E8F5E9")
        for r in row.cells[0].paragraphs[0].runs:
            r.font.color.rgb = WHITE; r.font.bold = True; r.font.size = Pt(10)
        for r in row.cells[1].paragraphs[0].runs:
            r.font.size = Pt(10)

    doc.add_page_break()

    # ── Introduccion ───────────────────────────────────────────────────────────
    h1(doc, "Diagramas de Casos de Uso")
    quote_block(doc, INTRO)
    doc.add_paragraph()

    # ── 1. Actores ─────────────────────────────────────────────────────────────
    h1(doc, "1. Especificacion de Actores")
    body(doc, (
        "El sistema Cantina Tita define cuatro actores principales. Cada actor representa "
        "un perfil de usuario con necesidades y niveles de acceso diferenciados. "
        "La jerarquia de roles garantiza que cada operacion sea ejecutada exclusivamente "
        "por el perfil autorizado, minimizando riesgos operativos y de seguridad."
    ))
    doc.add_paragraph()
    actor_table(doc, ACTORES)

    body(doc, (
        "El actor Sistema actua de forma automatizada como respuesta a eventos del negocio: "
        "cambios de stock, umbrales de saldo, cierres de caja y acciones de auditoria. "
        "No requiere intervencion humana directa y es el responsable de mantener la "
        "consistencia e integridad de los datos en todo momento."
    ))
    doc.add_page_break()

    # ── 2. Casos de uso por modulo ─────────────────────────────────────────────
    h1(doc, "2. Especificacion de Casos de Uso por Modulo")
    body(doc, (
        "Los casos de uso se agruparon por modulo, conservando la correspondencia uno a uno "
        "con los objetivos especificos del trabajo. Cada modulo detalla el objetivo asociado, "
        "los actores involucrados y la lista de casos de uso con su descripcion funcional."
    ))
    doc.add_paragraph()

    for mod in MODULOS:
        # cabecera de modulo
        h2(doc, f"Modulo {mod['num']}: {mod['titulo']}")

        label_value(doc, "Objetivo especifico", mod["objetivo"])
        label_value(doc, "Actores involucrados", mod["actores"])

        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("Casos de uso identificados:")
        r.font.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = NAVY
        doc.add_paragraph()

        cu_table(doc, mod["casos"])
        doc.add_page_break()

    # ── Resumen ────────────────────────────────────────────────────────────────
    h1(doc, "Resumen de Cobertura")
    body(doc, (
        "La siguiente tabla consolida la cantidad de casos de uso especificados por modulo "
        "y los actores que participan en cada uno, ofreciendo una vision global del alcance "
        "funcional del sistema."
    ))
    doc.add_paragraph()

    # Tabla resumen
    tbl_r = doc.add_table(rows=1, cols=4)
    tbl_r.style = "Table Grid"
    for i, txt in enumerate(["Modulo", "Nombre", "Actores principales", "Casos de uso"]):
        hdr = tbl_r.rows[0].cells[i]
        hdr.text = txt
        cell_bg(hdr, "1E3A5F")
        for r in hdr.paragraphs[0].runs:
            r.font.color.rgb = WHITE; r.font.bold = True; r.font.size = Pt(9)

    for ri, mod in enumerate(MODULOS):
        row = tbl_r.add_row().cells
        row[0].text = f"M-{mod['num']:0>2}"
        row[1].text = mod["titulo"]
        row[2].text = mod["actores"]
        row[3].text = str(len(mod["casos"]))
        bg = "E8F5E9" if ri % 2 == 0 else "FFFFFF"
        for ci in range(4):
            cell_bg(row[ci], bg)
            for r in row[ci].paragraphs[0].runs:
                r.font.size = Pt(9)
        for r in row[0].paragraphs[0].runs:
            r.font.bold = True; r.font.color.rgb = NAVY
        for r in row[1].paragraphs[0].runs:
            r.font.bold = True
        # cantidad en verde
        for r in row[3].paragraphs[0].runs:
            r.font.bold = True; r.font.color.rgb = GREEN

    doc.add_paragraph()

    total = sum(len(m["casos"]) for m in MODULOS)
    p = doc.add_paragraph()
    r1 = p.add_run("Total de casos de uso especificados: ")
    r1.font.size = Pt(11); r1.font.bold = True; r1.font.color.rgb = NAVY
    r2 = p.add_run(str(total))
    r2.font.size = Pt(14); r2.font.bold = True; r2.font.color.rgb = GREEN


# ══ MAIN ═══════════════════════════════════════════════════════════════════════
def main():
    doc = make_doc()
    build(doc)
    doc.save(str(OUT))
    print(f"Listo: {OUT}")
    total = sum(len(m["casos"]) for m in MODULOS)
    print(f"Modulos: {len(MODULOS)} | Casos de uso: {total} | Actores: {len(ACTORES)}")


if __name__ == "__main__":
    main()
